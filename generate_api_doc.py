from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Fonction pour styliser un paragraphe titre
def add_title_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    # For proper font on Windows:
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    p.space_after = Pt(6)
    return p

# Fonction pour un paragraphe normal
def add_normal_paragraph(doc, text):
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    # Fix font on Windows
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    p.space_after = Pt(4)
    return p

# Fonction pour ajouter un bloc de code avec fond gris clair
def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    shading_elm_1 = OxmlElement('w:shd')
    shading_elm_1.set(qn('w:fill'), 'F0F0F0')  # Gris clair
    p._p.get_or_add_pPr().append(shading_elm_1)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    p.space_after = Pt(8)
    # Optionally indent a bit
    p.paragraph_format.left_indent = Inches(0.25)
    return p

endpoints = [
    {
        "title": "/accounts/login (POST)",
        "description": "Permet à un utilisateur de se connecter avec identifiant et mot de passe, retourne un jeton JWT.",
        "addressed_to": "Client",
        "inputs_body": [
            ("identifiant", "Chaîne, Requis"),
            ("password", "Chaîne, Requis")
        ],
        "inputs_header": [],
        "example_body": '{\n  "identifiant": "123456",\n  "password": "123456"\n}',
        "example_header": None,
        "responses": [
            '200 : {"access_token": "..."}',
            '401 : {"message": "Identifiant ou mot de passe invalide"}',
            '500 : {"error": "Erreur interne du serveur"}'
        ]
    },
    {
        "title": "/accounts/logout (POST)",
        "description": "Permet à un utilisateur authentifié de se déconnecter en ajoutant son jeton JWT à une liste noire.",
        "addressed_to": "Client",
        "inputs_body": [],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": None,
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '200 : {"message": "Déconnexion réussie"}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/accounts/refer/<referral_code> (POST)",
        "description": "Permet à un nouvel utilisateur de s’inscrire avec un code de parrainage, crée un compte et attribue des récompenses.",
        "addressed_to": "Client",
        "inputs_body": [
            ("first_name", "Chaîne, Requis"),
            ("last_name", "Chaîne, Requis"),
            ("username", "Chaîne, Requis"),
            ("email", "Chaîne, Requis"),
            ("identifiant", "Chaîne, Requis"),
            ("password", "Chaîne, Requis")
        ],
        "inputs_header": [],
        "example_body": '{\n  "first_name": "John",\n  "last_name": "Doe",\n  "username": "johndoe",\n  "email": "john@example.com",\n  "identifiant": "123456",\n  "password": "password123",\n  "referral_code": "123e4567-e89b-12d3-a456-426614174000"\n}',
        "example_header": None,
        "responses": [
            '201 : {"message": "Inscription réussie via parrainage", "access_token": "..."}',
            '400 : {"message": "Code de parrainage invalide"}'
        ]
    },
    {
        "title": "/customer/invite (POST)",
        "description": "Permet à un utilisateur authentifié d’inviter un ami via email, crée un lien de parrainage.",
        "addressed_to": "Client",
        "inputs_body": [
            ("email", "Chaîne, Requis")
        ],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": '{\n  "email": "ami@example.com"\n}',
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '201 : {"message": "Invitation envoyée", "referral_link": "..."}',
            '400 : {"message": "Cet email a déjà été invité"}',
            '401 : {"msg": "En-tête d’autorisation manquant"}',
            '404 : {"message": "Utilisateur non trouvé"}'
        ]
    },
    {
        "title": "/lot/favorites (GET)",
        "description": "Permet à un utilisateur authentifié de récupérer sa liste de récompenses favorites.",
        "addressed_to": "Client",
        "inputs_body": [],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": None,
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '200 : {"favorites": [...]}',
            '401 : {"msg": "En-tête d’autorisation manquant"}',
            '404 : {"message": "Utilisateur non trouvé"}'
        ]
    },
    {
        "title": "/lot/favorites (POST)",
        "description": "Permet à un utilisateur authentifié d’ajouter une récompense à ses favoris.",
        "addressed_to": "Client",
        "inputs_body": [
            ("reward_id", "Entier, Requis")
        ],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": '{\n  "reward_id": 1\n}',
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '201 : {"message": "Récompense ajoutée aux favoris"}',
            '400 : {"message": "Récompense déjà dans les favoris"}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/lot/cart (GET)",
        "description": "Permet à un utilisateur authentifié de récupérer les articles de son panier.",
        "addressed_to": "Client",
        "inputs_body": [],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": None,
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '200 : {"cart_items": [...]}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/lot/cart (POST)",
        "description": "Permet à un utilisateur authentifié d’ajouter un article à son panier.",
        "addressed_to": "Client",
        "inputs_body": [
            ("reward_id", "Entier, Requis"),
            ("quantity", "Entier, Optionnel, défaut: 1")
        ],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": '{\n  "reward_id": 1,\n  "quantity": 2\n}',
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '201 : {"message": "Article ajouté au panier"}',
            '400 : {"message": "Récompense introuvable"}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/lot/orders (GET)",
        "description": "Permet à un utilisateur authentifié de récupérer la liste de ses commandes.",
        "addressed_to": "Client",
        "inputs_body": [],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": None,
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '200 : {"orders": [...]}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/lot/orders (POST)",
        "description": "Permet à un utilisateur authentifié de passer une commande à partir de son panier.",
        "addressed_to": "Client",
        "inputs_body": [
            ("contact", "Chaîne, Optionnel")
        ],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": '{\n  "contact": "john@example.com"\n}',
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '201 : {"message": "Commande passée avec succès"}',
            '400 : {"message": "Solde insuffisant"}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/lot/notifications (GET)",
        "description": "Permet à un utilisateur authentifié de récupérer ses notifications.",
        "addressed_to": "Client",
        "inputs_body": [],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": None,
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '200 : {"notifications": [...]}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/faq (GET)",
        "description": "Permet à un utilisateur de récupérer la liste des FAQ.",
        "addressed_to": "Client",
        "inputs_body": [],
        "inputs_header": [],
        "example_body": None,
        "example_header": None,
        "responses": [
            '200 : {"faqs": [...]}'
        ]
    },
    {
        "title": "/support (GET)",
        "description": "Permet à un utilisateur authentifié de récupérer ses requêtes de support.",
        "addressed_to": "Client",
        "inputs_body": [],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": None,
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '200 : {"support_requests": [...]}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/support (POST)",
        "description": "Permet à un utilisateur authentifié de créer une nouvelle requête de support.",
        "addressed_to": "Client",
        "inputs_body": [
            ("subject", "Chaîne, Requis"),
            ("description", "Chaîne, Requis"),
            ("request_type", 'Chaîne, Requis : "Réclamation", "Assistance", "Autre"')
        ],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": '{\n  "subject": "Problème de paiement",\n  "description": "Facturé deux fois",\n  "request_type": "Réclamation"\n}',
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '201 : {"message": "Requête de support créée"}',
            '400 : {"message": "Le sujet est requis"}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/survey (GET)",
        "description": "Permet à un utilisateur de récupérer la liste des enquêtes actives.",
        "addressed_to": "Client",
        "inputs_body": [],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": None,
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '200 : {"surveys": [...]}',
            '401 : {"msg": "En-tête d’autorisation manquant"}'
        ]
    },
    {
        "title": "/survey/<survey_id>/respond (POST)",
        "description": "Permet à un utilisateur authentifié de soumettre une réponse à une enquête.",
        "addressed_to": "Client",
        "inputs_body": [
            ("option_id", "Entier, Requis")
        ],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": '{\n  "option_id": 1\n}',
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '201 : {"message": "Réponse soumise"}',
            '400 : {"message": "Vous avez déjà répondu à cette enquête"}',
            '401 : {"msg": "En-tête d’autorisation manquant"}',
            '404 : {"message": "Enquête non trouvée"}'
        ]
    },
    {
        "title": "/dashboard (GET)",
        "description": "Permet à un utilisateur authentifié de récupérer les informations de son tableau de bord (solde, commandes, notifications).",
        "addressed_to": "Client",
        "inputs_body": [],
        "inputs_header": [
            ("Authorization", "Bearer <token>")
        ],
        "example_body": None,
        "example_header": "Authorization: Bearer eyJ...",
        "responses": [
            '200 : {"balance": 500, "recent_orders": [...], "notifications": [...]}',
            '401 : {"msg": "En-tête d’autorisation manquant"}',
            '404 : {"message": "Utilisateur non trouvé"}'
        ]
    }
]


def create_document(filename='Api_Documentation_MyWitti.docx'):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    doc.add_paragraph("Documentation API MyWitti", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    for i, ep in enumerate(endpoints, start=1):
        # Titre
        add_title_paragraph(doc, f"{i}. {ep['title']}")

        # Description
        add_normal_paragraph(doc, f"- Description : {ep['description']}")
        add_normal_paragraph(doc, f"- Adressé à : {ep['addressed_to']}")

        # Entrées header
        if ep["inputs_header"]:
            add_normal_paragraph(doc, "- Entrées nécessaires (header) :")
            for name, desc in ep["inputs_header"]:
                add_normal_paragraph(doc, f"  - {name}: {desc}")

        # Entrées body
        if ep["inputs_body"]:
            add_normal_paragraph(doc, "- Entrées nécessaires (body) :")
            for name, desc in ep["inputs_body"]:
                add_normal_paragraph(doc, f"  - {name}: {desc}")

        # Exemple header
        if ep["example_header"]:
            add_normal_paragraph(doc, "- Exemple de test (header) :")
            add_code_block(doc, ep["example_header"])

        # Exemple body
        if ep["example_body"]:
            add_normal_paragraph(doc, "- Exemple de test (body) :")
            add_code_block(doc, ep["example_body"])

        # Réponses possibles
        add_normal_paragraph(doc, "- Réponses possibles :")
        for resp in ep["responses"]:
            add_normal_paragraph(doc, f"  - {resp}")

        doc.add_paragraph()  # saut de ligne

    doc.save(filename)
    print(f"Document généré : {filename}")

if __name__ == "__main__":
    # Pour générer le doc avec uniquement les deux exemples d'endpoints actuels
    create_document()
